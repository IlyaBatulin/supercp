from pathlib import Path

from docx import Document


def add_paragraphs(doc: Document, text: str) -> None:
    for block in text.strip().split("\n\n"):
        doc.add_paragraph(block.strip())


def main() -> None:
    out_dir = Path("docs") / "LR4"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_file = out_dir / "LR4_report_clean.docx"

    doc = Document()
    doc.add_heading("Лабораторная работа №4", level=1)
    doc.add_paragraph(
        "Тема: Управление требованиями при развертывании приложений в SaaS Replit на примере фреймворка Django."
    )

    doc.add_heading("Введение", level=2)
    add_paragraphs(
        doc,
        """
В данной лабораторной работе я изучил управление требованиями при развертывании веб-приложения в SaaS-среде Replit на базе Django. Цель работы состояла в том, чтобы развернуть приложение, оформить структуру проекта по требованиям задания и подготовить рабочую конфигурацию запуска.

В качестве практической основы я использовал компонент контроля доступа к вычислительному кластеру: вход по логину, паролю и ключу, выбор узла, проверка пароля узла и фиксация статистики сессии. Дополнительно я доработал интерфейс и оформил описание параметров запуска.
""",
    )

    doc.add_heading("Алгоритм выполнения работы", level=2)
    steps = [
        "1. Я проверил структуру Django-проекта: наличие manage.py, settings.py, urls.py и зависимостей в requirements.txt.",
        "2. Я создал отдельное приложение tasks с файлами apps.py, views.py, urls.py и шаблоном tasks/templates/tasks/index.html.",
        "3. Я подключил приложение tasks в INSTALLED_APPS и добавил маршрутизацию в корневой urls.py через include('tasks.urls').",
        "4. Я реализовал представление index(), которое обрабатывает вход, выбор узла и выход, используя основную логику из main.py.",
        "5. Я оформил index.html с комментариями, добавил раздел Description и блок «Этапы задач на сегодня» (не менее 7 задач).",
        "6. Я настроил переменные окружения в .env.example, включая STUDENT_FIO и параметры Django/Replit.",
        "7. Я подготовил скрипт scripts/deploy_replit.sh: установка зависимостей, миграции, запуск сервера на PORT.",
        "8. Я выполнил проверку конфигурации командой python manage.py check и убедился, что системных ошибок нет.",
        "9. Я улучшил дизайн страницы tasks: современная тема, карточный интерфейс, акцентные кнопки и адаптивная верстка.",
    ]
    for step in steps:
        doc.add_paragraph(step)

    doc.add_heading("Ключевые фрагменты кода", level=2)
    doc.add_paragraph("Подключение приложения tasks в settings.py:")
    doc.add_paragraph(
        'INSTALLED_APPS = [ ..., "tasks", "access_control" ]',
        style="Intense Quote",
    )
    doc.add_paragraph("Маршрутизация в корневом urls.py:")
    doc.add_paragraph(
        'urlpatterns = [ path("", include("tasks.urls")), path("access-control/", include("access_control.urls")) ]',
        style="Intense Quote",
    )
    doc.add_paragraph("Запуск в Replit через deploy-скрипт:")
    doc.add_paragraph(
        "pip install -r requirements.txt\npython manage.py migrate\npython manage.py runserver \"0.0.0.0:${PORT:-8000}\"",
        style="Intense Quote",
    )

    doc.add_heading("Что реализовано по итогу", level=2)
    add_paragraphs(
        doc,
        """
По итогам работы у меня есть готовое Django-приложение tasks, подключенное как основной маршрут проекта. Страница index.html формируется через представление, обрабатывает пользовательский сценарий и содержит комментарии, описание переменных и этапы задач.

Запуск под Replit автоматизирован через .replit и скрипт deploy_replit.sh. Конфигурация проверена, приложение запускается и корректно обрабатывает вход, выбор узла и выдачу результата.
""",
    )

    doc.add_heading("Заключение", level=2)
    add_paragraphs(
        doc,
        """
В результате выполнения лабораторной работы я закрепил навыки настройки Django-проекта, подключения приложения, маршрутизации и подготовки конфигурации для SaaS Replit. Я также оформил интерфейс и документационную часть так, чтобы проект был удобен для демонстрации и повторного развёртывания.

Полученный результат соответствует основной цели лабораторной работы: управление требованиями к веб-компоненту и его развертыванию в облачной учебной среде.
""",
    )

    doc.save(out_file)


if __name__ == "__main__":
    main()
